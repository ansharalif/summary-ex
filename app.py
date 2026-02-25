import streamlit as st
import pandas as pd
from io import BytesIO
from datetime import datetime

# =========================================================
# SAFE IMPORTS (agar app tidak langsung crash)
# =========================================================
OPENPYXL_AVAILABLE = True
DOCX_AVAILABLE = True
IMPORT_ERRORS = []

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ModuleNotFoundError as e:
    OPENPYXL_AVAILABLE = False
    IMPORT_ERRORS.append(str(e))

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ModuleNotFoundError as e:
    DOCX_AVAILABLE = False
    IMPORT_ERRORS.append(str(e))


# =========================================================
# KONFIGURASI APP
# =========================================================
st.set_page_config(
    page_title="Policy Matrix Ringkas Generator",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Template “Policy Matrix Ringkas” – Analisis Cepat Pimpinan")
st.caption("Generator template Excel + contoh negara + brief Word 1 halaman (versi Streamlit).")

if IMPORT_ERRORS:
    st.warning(
        "Beberapa dependensi belum terpasang di environment ini. "
        "Fitur tertentu mungkin tidak aktif.\n\n"
        "Error terdeteksi:\n- " + "\n- ".join(IMPORT_ERRORS)
    )


# =========================================================
# DATA CONTOH
# =========================================================
SAMPLE_DATA = [
    {
        "Negara": "Indonesia",
        "Last update": "2026-02-20",
        "Pendekatan dominan": "Kombinasi penegakan hukum, pencegahan, dan rehabilitasi",
        "Instrumen hukum utama": "Peraturan perundang-undangan terkait tindak pidana terorisme dan kebijakan turunan",
        "Catatan kehati-hatian (HAM/rule of law)": "Perlu menjaga keseimbangan antara keamanan, due process, dan perlindungan hak-hak dasar",
        "Relevansi pembelajaran untuk Indonesia/BNPT": "Penguatan standardisasi pemetaan kebijakan dan monitoring pembaruan lintas isu",
        "Sumber ringkas": "https://www.foreignterroristfighters.info/policy-matrix",
    },
    {
        "Negara": "France",
        "Last update": "2026-02-15",
        "Pendekatan dominan": "Penegakan hukum dan pengawasan yang kuat, disertai langkah pencegahan",
        "Instrumen hukum utama": "Kerangka hukum nasional kontra-terorisme dan mekanisme proses pidana",
        "Catatan kehati-hatian (HAM/rule of law)": "Perlu perhatian pada proporsionalitas pengawasan dan perlindungan kebebasan sipil",
        "Relevansi pembelajaran untuk Indonesia/BNPT": "Pembelajaran komparatif untuk aspek koordinasi penegakan dan pengawasan berbasis hukum",
        "Sumber ringkas": "https://www.foreignterroristfighters.info/policy-matrix",
    },
    {
        "Negara": "Germany",
        "Last update": "2026-02-12",
        "Pendekatan dominan": "Kombinasi penegakan, pencegahan, dan program rehabilitasi/reintegrasi",
        "Instrumen hukum utama": "Ketentuan pidana, pengawasan, dan kebijakan reintegrasi berbasis kelembagaan",
        "Catatan kehati-hatian (HAM/rule of law)": "Penting memastikan non-diskriminasi, due process, dan evaluasi efektivitas program",
        "Relevansi pembelajaran untuk Indonesia/BNPT": "Referensi untuk pengembangan kerangka pembelajaran rehabilitasi dan koordinasi multi-aktor",
        "Sumber ringkas": "https://www.foreignterroristfighters.info/policy-matrix",
    },
    {
        "Negara": "Kazakhstan",
        "Last update": "2026-02-05",
        "Pendekatan dominan": "Kombinasi repatriasi, reintegrasi, dan langkah keamanan",
        "Instrumen hukum utama": "Kebijakan nasional penanganan returnees dan dukungan kelembagaan terkait",
        "Catatan kehati-hatian (HAM/rule of law)": "Perlu perhatian pada perlindungan kelompok rentan dan akuntabilitas implementasi",
        "Relevansi pembelajaran untuk Indonesia/BNPT": "Bahan belajar untuk isu repatriasi, reintegrasi, dan dukungan keluarga/perempuan-anak",
        "Sumber ringkas": "https://www.foreignterroristfighters.info/policy-matrix",
    },
    {
        "Negara": "Philippines",
        "Last update": "2026-01-28",
        "Pendekatan dominan": "Penegakan hukum dengan dukungan pencegahan; data reintegrasi perlu diperdalam",
        "Instrumen hukum utama": "Perangkat hukum kontra-terorisme nasional dan kebijakan operasional terkait",
        "Catatan kehati-hatian (HAM/rule of law)": "Perlu menjaga akuntabilitas, pengawasan, dan perlindungan HAM dalam implementasi",
        "Relevansi pembelajaran untuk Indonesia/BNPT": "Pembelajaran regional untuk komparasi pendekatan dan identifikasi gap data kebijakan",
        "Sumber ringkas": "https://www.foreignterroristfighters.info/policy-matrix",
    },
]

COLUMNS = [
    "Negara",
    "Last update",
    "Pendekatan dominan",
    "Instrumen hukum utama",
    "Catatan kehati-hatian (HAM/rule of law)",
    "Relevansi pembelajaran untuk Indonesia/BNPT",
    "Sumber ringkas",
]

df = pd.DataFrame(SAMPLE_DATA)


# =========================================================
# HELPER (EXCEL)
# =========================================================
def apply_thin_border(cell):
    thin = Side(style="thin", color="D9D9D9")
    cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)

def style_header_row(ws, row_idx: int, start_col: int, end_col: int):
    fill = PatternFill(fill_type="solid", fgColor="1F4E78")
    font = Font(color="FFFFFF", bold=True)
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for col in range(start_col, end_col + 1):
        cell = ws.cell(row=row_idx, column=col)
        cell.fill = fill
        cell.font = font
        cell.alignment = align
        apply_thin_border(cell)

def auto_width(ws, min_width=12, max_width=55):
    for col_cells in ws.columns:
        col_letter = get_column_letter(col_cells[0].column)
        max_len = 0
        for c in col_cells:
            val = "" if c.value is None else str(c.value)
            max_len = max(max_len, len(val))
        ws.column_dimensions[col_letter].width = min(max(max_len + 2, min_width), max_width)

def build_excel_workbook_bytes(sample_data):
    if not OPENPYXL_AVAILABLE:
        raise ModuleNotFoundError("openpyxl belum terpasang")

    wb = Workbook()

    # Sheet 1: Template
    ws_template = wb.active
    ws_template.title = "Template"
    ws_template["A1"] = "Template Policy Matrix Ringkas - Analisis Cepat Pimpinan"
    ws_template["A1"].font = Font(size=13, bold=True)
    ws_template.merge_cells("A1:G1")

    ws_template["A2"] = "Isi 1 baris = 1 country brief ringkas (1 halaman saat diekspor ke Word)"
    ws_template["A2"].font = Font(italic=True, color="666666")
    ws_template.merge_cells("A2:G2")

    for i, c in enumerate(COLUMNS, start=1):
        ws_template.cell(row=4, column=i, value=c)
    style_header_row(ws_template, 4, 1, len(COLUMNS))

    for r in range(5, 10):
        for c in range(1, len(COLUMNS) + 1):
            val = "YYYY-MM-DD" if c == 2 else ""
            cell = ws_template.cell(row=r, column=c, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            apply_thin_border(cell)

    ws_template.freeze_panes = "A5"

    # Sheet 2: Contoh
    ws_sample = wb.create_sheet("Contoh_3-5_Negara")
    ws_sample["A1"] = "Contoh Isian Policy Matrix Ringkas (5 Negara)"
    ws_sample["A1"].font = Font(size=13, bold=True)
    ws_sample.merge_cells("A1:G1")

    for i, c in enumerate(COLUMNS, start=1):
        ws_sample.cell(row=3, column=i, value=c)
    style_header_row(ws_sample, 3, 1, len(COLUMNS))

    for r_idx, item in enumerate(sample_data, start=4):
        for c_idx, col_name in enumerate(COLUMNS, start=1):
            cell = ws_sample.cell(row=r_idx, column=c_idx, value=item.get(col_name, ""))
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            apply_thin_border(cell)

    ws_sample.freeze_panes = "A4"

    # Sheet 3: Panduan
    ws_guide = wb.create_sheet("Panduan_Pengisian")
    ws_guide["A1"] = "Panduan Pengisian Singkat - Policy Matrix Ringkas"
    ws_guide["A1"].font = Font(size=13, bold=True)
    ws_guide.merge_cells("A1:D1")

    guide_rows = [
        ["Kolom", "Isi", "Contoh", "Catatan"],
        ["Negara", "Nama negara", "France", "1 baris untuk 1 negara / 1 isu"],
        ["Last update", "Tanggal pembaruan terakhir", "2026-02-15", "Gunakan format YYYY-MM-DD"],
        ["Pendekatan dominan", "Ringkasan pendekatan utama", "Penegakan + pencegahan", "Tulis singkat, 1-2 kalimat"],
        ["Instrumen hukum utama", "Instrumen/kerangka hukum kunci", "Ketentuan pidana + kebijakan turunan", "Tidak perlu terlalu detail"],
        ["Catatan kehati-hatian (HAM/rule of law)", "Poin kehati-hatian implementasi", "Due process, proporsionalitas", "Gunakan bahasa netral dan analitis"],
        ["Relevansi pembelajaran untuk Indonesia/BNPT", "Nilai pembelajaran untuk internal", "Komparasi, lesson learned", "Fokus pembelajaran, bukan evaluasi normatif"],
        ["Sumber ringkas", "URL sumber utama", "https://www.foreignterroristfighters.info/policy-matrix", "Cantumkan sumber untuk penelusuran"],
    ]

    for r_idx, row in enumerate(guide_rows, start=3):
        for c_idx, val in enumerate(row, start=1):
            cell = ws_guide.cell(row=r_idx, column=c_idx, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            apply_thin_border(cell)
    style_header_row(ws_guide, 3, 1, 4)

    for ws in [ws_template, ws_sample, ws_guide]:
        auto_width(ws)
        for r in range(1, ws.max_row + 1):
            ws.row_dimensions[r].height = 22

    for ws in [ws_template, ws_sample]:
        ws.column_dimensions["C"].width = 38
        ws.column_dimensions["D"].width = 38
        ws.column_dimensions["E"].width = 45
        ws.column_dimensions["F"].width = 45
        ws.column_dimensions["G"].width = 42

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output


# =========================================================
# HELPER (DOCX)
# =========================================================
def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_doc_margins(document, top=0.6, bottom=0.6, left=0.7, right=0.7):
    section = document.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def build_docx_bytes(item: dict):
    if not DOCX_AVAILABLE:
        raise ModuleNotFoundError("python-docx belum terpasang")

    doc = Document()
    set_doc_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("POLICY MATRIX RINGKAS - ANALISIS CEPAT PIMPINAN")
    run.bold = True
    run.font.size = Pt(13)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run(f"Negara: {item.get('Negara', '-')}")
    subrun.italic = True
    subrun.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    rows_content = [
        ("Negara", item.get("Negara", "-")),
        ("Last update", item.get("Last update", "-")),
        ("Pendekatan dominan", item.get("Pendekatan dominan", "-")),
        ("Instrumen hukum utama", item.get("Instrumen hukum utama", "-")),
        ("Catatan kehati-hatian (HAM/rule of law)", item.get("Catatan kehati-hatian (HAM/rule of law)", "-")),
        ("Relevansi pembelajaran untuk Indonesia/BNPT", item.get("Relevansi pembelajaran untuk Indonesia/BNPT", "-")),
        ("Sumber ringkas", item.get("Sumber ringkas", "-")),
    ]

    for label, value in rows_content:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].width = Inches(2.2)
        cells[1].width = Inches(4.9)

        for i, c in enumerate(cells):
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(10)
                    if i == 0:
                        r.bold = True
            if i == 0:
                set_cell_shading(c, "D9E1F2")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        "Catatan: Ringkasan internal untuk analisis cepat. "
        "Lakukan verifikasi lanjutan pada sumber primer sebelum dipakai untuk keputusan substantif."
    )
    foot_run.font.size = Pt(9)
    foot_run.italic = True

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =========================================================
# UI
# =========================================================
left, right = st.columns([1.2, 1])

with left:
    st.subheader("Preview Data Contoh (5 Negara)")
    st.dataframe(df, use_container_width=True, hide_index=True)

with right:
    st.subheader("Info Fitur")
    st.markdown(
        """
- **Excel Template**: Template + Contoh 5 negara + Panduan pengisian  
- **Word Brief**: 1 halaman per negara (siap untuk pimpinan)  
- **Mode Aman**: App tetap tampil walau package belum terpasang  
        """
    )
    st.info(f"Waktu akses: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")


st.divider()
st.subheader("⬇️ Download Template & Brief")

col1, col2 = st.columns(2)

with col1:
    st.markdown("### 1) Template Excel")
    if OPENPYXL_AVAILABLE:
        try:
            excel_bytes = build_excel_workbook_bytes(SAMPLE_DATA)
            st.download_button(
                label="Download Template_Policy_Matrix_Ringkas.xlsx",
                data=excel_bytes,
                file_name="Template_Policy_Matrix_Ringkas.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            st.success("Template Excel siap diunduh.")
        except Exception as e:
            st.error(f"Gagal membuat Excel: {e}")
    else:
        st.error("Fitur Excel nonaktif: package `openpyxl` belum terpasang.")

with col2:
    st.markdown("### 2) Word Brief per Negara")
    negara_opsi = [item["Negara"] for item in SAMPLE_DATA]
    selected_negara = st.selectbox("Pilih negara", negara_opsi)

    if DOCX_AVAILABLE:
        selected_item = next((x for x in SAMPLE_DATA if x["Negara"] == selected_negara), None)
        if selected_item:
            try:
                docx_bytes = build_docx_bytes(selected_item)
                filename = f"Policy_Matrix_Ringkas_{selected_negara.replace(' ', '_')}.docx"
                st.download_button(
                    label=f"Download {filename}",
                    data=docx_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                st.success("Brief Word siap diunduh.")
            except Exception as e:
                st.error(f"Gagal membuat Word brief: {e}")
    else:
        st.error("Fitur Word nonaktif: package `python-docx` belum terpasang.")


st.divider()
st.subheader("🛠️ Petunjuk Deploy (Streamlit Cloud)")
st.code(
    """# requirements.txt
streamlit>=1.30
pandas>=2.0
openpyxl>=3.1.2
python-docx>=1.1.0
""",
    language="txt"
)

st.caption(
    "Jika app sebelumnya error ModuleNotFoundError, biasanya cukup tambahkan package ke requirements.txt lalu redeploy."
)
